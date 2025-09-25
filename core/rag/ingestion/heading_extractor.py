import re
import os
from typing import Dict, List, Any, Optional
from docx import Document
import fitz  # PyMuPDF

class HeadingExtractor:
    """Extract headings dynamically from documents"""
    
    def __init__(self):
        self.heading_keywords = [
            'summary', 'introduction', 'overview', 'background', 'scope', 'methodology',
            'findings', 'results', 'conclusion', 'recommendation', 'risk', 'mitigation',
            'objective', 'goal', 'purpose', 'approach', 'analysis', 'assessment'
        ]
    
    def extract_sections(self, file_path: str) -> Dict[str, str]:
        """
        Extract sections from document based on detected headings.
        
        Args:
            file_path: Path to document file
            
        Returns:
            Dict mapping section names to their content
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.docx':
            return self._extract_docx_sections(file_path)
        elif file_ext == '.pdf':
            return self._extract_pdf_sections(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
    
    def extract_docx_headings(self, doc: Document) -> List[Dict[str, Any]]:
        """Extract headings from DOCX document"""
        headings = []
        current_section = None
        current_content = []
        section_counter = 1
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            
            # Check if this paragraph is a heading
            if self._is_docx_heading(paragraph):
                # Save previous section if exists
                if current_section and current_content:
                    headings.append({
                        'section_id': f'sec_{section_counter:03d}',
                        'heading': current_section,
                        'content': '\n'.join(current_content).strip(),
                        'level': self._get_docx_heading_level(paragraph),
                        'page_start': 1,  # DOCX doesn't have clear page concept
                        'page_end': 1
                    })
                    section_counter += 1
                
                # Start new section
                current_section = text
                current_content = []
            else:
                # Add to current section content
                if current_section:
                    current_content.append(text)
                else:
                    # Content before first heading - create introduction
                    if not headings or headings[0]['heading'] != 'Introduction':
                        headings.insert(0, {
                            'section_id': 'sec_000',
                            'heading': 'Introduction',
                            'content': text,
                            'level': 1,
                            'page_start': 1,
                            'page_end': 1
                        })
                    else:
                        headings[0]['content'] += '\n' + text
        
        # Save final section
        if current_section and current_content:
            headings.append({
                'section_id': f'sec_{section_counter:03d}',
                'heading': current_section,
                'content': '\n'.join(current_content).strip(),
                'level': self._get_docx_heading_level(paragraph),
                'page_start': 1,
                'page_end': 1
            })
        
        return headings
    
    def extract_pdf_headings(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract headings from PDF document"""
        try:
            doc = fitz.open(file_path)
            headings = []
            current_section = None
            current_content = []
            section_counter = 1
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                blocks = page.get_text("dict")["blocks"]
                
                for block in blocks:
                    if "lines" in block:
                        for line in block["lines"]:
                            for span in line["spans"]:
                                text = span["text"].strip()
                                if not text:
                                    continue
                                
                                # Check if this looks like a heading
                                if self._is_pdf_heading(span, text):
                                    # Save previous section
                                    if current_section and current_content:
                                        headings.append({
                                            'section_id': f'sec_{section_counter:03d}',
                                            'heading': current_section,
                                            'content': ' '.join(current_content).strip(),
                                            'level': self._get_pdf_heading_level(span),
                                            'page_start': page_num + 1,
                                            'page_end': page_num + 1
                                        })
                                        section_counter += 1
                                    
                                    # Start new section
                                    current_section = text
                                    current_content = []
                                else:
                                    # Add to current section content
                                    if current_section:
                                        current_content.append(text)
                                    else:
                                        # Content before first heading
                                        if not headings or headings[0]['heading'] != 'Introduction':
                                            headings.insert(0, {
                                                'section_id': 'sec_000',
                                                'heading': 'Introduction',
                                                'content': text,
                                                'level': 1,
                                                'page_start': page_num + 1,
                                                'page_end': page_num + 1
                                            })
                                        else:
                                            headings[0]['content'] += ' ' + text
            
            # Save final section
            if current_section and current_content:
                headings.append({
                    'section_id': f'sec_{section_counter:03d}',
                    'heading': current_section,
                    'content': ' '.join(current_content).strip(),
                    'level': 1,
                    'page_start': len(doc),
                    'page_end': len(doc)
                })
            
            doc.close()
            return headings
            
        except Exception as e:
            print(f"Error extracting PDF headings: {e}")
            return []
    
    def _extract_docx_sections(self, file_path: str) -> Dict[str, str]:
        """Extract sections from DOCX file"""
        try:
            doc = Document(file_path)
            
            sections = {}
            current_heading = None
            current_content = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                
                # Check if this paragraph is a heading
                if self._is_docx_heading(paragraph):
                    # Save previous section if exists
                    if current_heading and current_content:
                        sections[current_heading] = '\n'.join(current_content).strip()
                    
                    # Start new section
                    current_heading = text
                    current_content = []
                else:
                    # Add to current section content
                    if current_heading:
                        current_content.append(text)
                    else:
                        # Content before first heading
                        if "Introduction" not in sections:
                            sections["Introduction"] = ""
                        sections["Introduction"] += text + "\n"
            
            # Save final section
            if current_heading and current_content:
                sections[current_heading] = '\n'.join(current_content).strip()
            
            # Fallback if no sections found
            if not sections:
                full_text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
                sections["Document Content"] = full_text
            
            return sections
            
        except Exception as e:
            print(f"Error extracting DOCX sections: {e}")
            return {"Document Content": "Error extracting content"}
    
    def _extract_pdf_sections(self, file_path: str) -> Dict[str, str]:
        """Extract sections from PDF file"""
        try:
            doc = fitz.open(file_path)
            
            sections = {}
            current_heading = None
            current_content = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                blocks = page.get_text("dict")["blocks"]
                
                for block in blocks:
                    if "lines" in block:
                        for line in block["lines"]:
                            for span in line["spans"]:
                                text = span["text"].strip()
                                if not text:
                                    continue
                                
                                # Check if this looks like a heading
                                if self._is_pdf_heading(span, text):
                                    # Save previous section
                                    if current_heading and current_content:
                                        sections[current_heading] = ' '.join(current_content).strip()
                                    
                                    # Start new section
                                    current_heading = text
                                    current_content = []
                                else:
                                    # Add to current section content
                                    if current_heading:
                                        current_content.append(text)
                                    else:
                                        # Content before first heading
                                        if "Introduction" not in sections:
                                            sections["Introduction"] = ""
                                        sections["Introduction"] += text + " "
            
            # Save final section
            if current_heading and current_content:
                sections[current_heading] = ' '.join(current_content).strip()
            
            # Fallback if no sections found
            if not sections:
                full_text = ""
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    full_text += page.get_text() + "\n"
                sections["Document Content"] = full_text.strip()
            
            doc.close()
            return sections
            
        except Exception as e:
            print(f"Error extracting PDF sections: {e}")
            return {"Document Content": "Error extracting content"}
    
    def _is_docx_heading(self, paragraph) -> bool:
        """Check if DOCX paragraph is a heading"""
        # Check if it's a built-in heading style
        if paragraph.style.name.startswith('Heading'):
            return True
        
        # Check if it's bold and short (likely a heading)
        if paragraph.runs:
            first_run = paragraph.runs[0]
            text = paragraph.text.strip()
            if (first_run.bold and 
                len(text) < 100 and 
                not text.endswith('.') and
                len(text.split()) <= 10):
                return True
        
        # Check for numbering patterns
        text = paragraph.text.strip()
        if re.match(r'^\d+\.?\s+[A-Z]', text):  # "1. Title" or "1 Title"
            return True
        if re.match(r'^\d+\.\d+\.?\s+[A-Z]', text):  # "1.1 Title"
            return True
        if re.match(r'^[A-Z]\.?\s+[A-Z]', text):  # "A. Title"
            return True
        
        return False
    
    def _is_pdf_heading(self, span: dict, text: str) -> bool:
        """Check if PDF span is a heading"""
        # Check font size (headings are usually larger)
        font_size = span.get('size', 12)
        if font_size > 14:
            return True
        
        # Check if bold and short
        font_flags = span.get('flags', 0)
        is_bold = bool(font_flags & 2**4)  # Bold flag
        
        if (is_bold and 
            len(text) < 100 and 
            not text.endswith('.') and
            len(text.split()) <= 10):
            return True
        
        # Check for numbering patterns
        if re.match(r'^\d+\.?\s+[A-Z]', text):  # "1. Title"
            return True
        if re.match(r'^\d+\.\d+\.?\s+[A-Z]', text):  # "1.1 Title"
            return True
        if re.match(r'^[A-Z]\.?\s+[A-Z]', text):  # "A. Title"
            return True
        
        return False
    
    def _get_docx_heading_level(self, paragraph) -> int:
        """Get heading level from DOCX paragraph"""
        style_name = paragraph.style.name
        if 'Heading' in style_name:
            try:
                return int(style_name.split()[-1])
            except:
                return 1
        return 1
    
    def _get_pdf_heading_level(self, span: dict) -> int:
        """Get heading level from PDF span based on font size"""
        font_size = span.get('size', 12)
        if font_size >= 18:
            return 1
        elif font_size >= 16:
            return 2
        elif font_size >= 14:
            return 3
        else:
            return 1
    
    def map_headings_to_sections(self, detected_headings: List[Dict], target_sections: List[str]) -> Dict[str, Dict]:
        """
        Map detected headings to target report sections using semantic similarity.
        
        Args:
            detected_headings: List of detected heading dictionaries
            target_sections: List of target section names (e.g., ['executive_summary', 'scope_of_work'])
            
        Returns:
            Dict mapping target sections to best matching detected headings
        """
        mapping = {}
        
        # Define mapping rules for common sections
        section_keywords = {
            'executive_summary': ['summary', 'executive', 'overview', 'introduction', 'abstract'],
            'scope_of_work': ['scope', 'work', 'methodology', 'approach', 'method', 'process'],
            'value_proposition': ['value', 'proposition', 'benefit', 'advantage', 'strength'],
            'why_company_a': ['why', 'company', 'about', 'background', 'profile', 'experience'],
            'risks': ['risk', 'challenge', 'issue', 'concern', 'mitigation', 'problem'],
            'recommendations': ['recommendation', 'conclusion', 'next', 'steps', 'action', 'future'],
            'findings': ['finding', 'result', 'outcome', 'discovery', 'analysis', 'conclusion']
        }
        
        for target_section in target_sections:
            best_match = None
            best_score = 0
            
            # Get keywords for this target section
            keywords = section_keywords.get(target_section, [target_section.replace('_', ' ').split()])
            
            for heading_data in detected_headings:
                heading_text = heading_data['heading'].lower()
                
                # Calculate similarity score
                score = 0
                for keyword in keywords:
                    if keyword in heading_text:
                        score += 1
                
                # Boost score for exact matches
                if any(keyword == heading_text for keyword in keywords):
                    score += 5
                
                if score > best_score:
                    best_score = score
                    best_match = heading_data
            
            if best_match:
                mapping[target_section] = best_match
            else:
                # Fallback: create a summarized section from all content
                all_content = ' '.join([h['content'] for h in detected_headings])
                mapping[target_section] = {
                    'heading': f"Summarized {target_section.replace('_', ' ').title()}",
                    'content': all_content[:1000] + "..." if len(all_content) > 1000 else all_content,
                    'level': 1,
                    'page_start': 1,
                    'page_end': 1,
                    'section_id': f'fallback_{target_section}'
                }
        
        return mapping
    
    def generate_section_schema(self, sections: Dict[str, str]) -> Dict[str, Any]:
        """
        Generate a dynamic schema from extracted sections.
        Each section gets default instructions for processing.
        
        Args:
            sections: Dict of section_name -> content
            
        Returns:
            Schema dict compatible with existing pipeline
        """
        schema = {}
        
        for section_name, content in sections.items():
            # Clean section name for use as key
            section_key = section_name.lower().replace(' ', '_').replace('.', '').replace(':', '')
            
            # Determine appropriate instructions based on section name
            instructions = self._get_section_instructions(section_name)
            
            schema[section_key] = {
                "title": section_name,
                "source": "dynamic",  # Indicates this was auto-detected
                "instructions": instructions,
                "content": ""  # Will be filled by drafting agent
            }
        
        return schema
    
    def _get_section_instructions(self, section_name: str) -> Dict[str, str]:
        """Get appropriate instructions based on section name"""
        section_lower = section_name.lower()
        
        # Executive summary type sections
        if any(word in section_lower for word in ['summary', 'executive', 'overview', 'introduction']):
            return {
                "objective": "Provide a clear and concise summary of this section, highlighting the key points and main ideas.",
                "tone": "Professional and informative",
                "length": "2-3 paragraphs, 150-200 words",
                "format": "Paragraph form with clear, flowing sentences"
            }
        
        # Scope/methodology sections
        elif any(word in section_lower for word in ['scope', 'methodology', 'approach', 'method']):
            return {
                "objective": "Clearly outline the scope, methodology, or approach described in this section with specific details and steps.",
                "tone": "Technical and precise",
                "length": "2-4 paragraphs, 200-300 words",
                "format": "Structured paragraphs with numbered or bulleted sub-points where appropriate"
            }
        
        # Risk/challenge sections
        elif any(word in section_lower for word in ['risk', 'challenge', 'issue', 'problem']):
            return {
                "objective": "Identify and explain the risks, challenges, or issues presented, along with any mitigation strategies mentioned.",
                "tone": "Analytical and solution-focused",
                "length": "2-3 paragraphs, 150-250 words",
                "format": "Clear identification of issues followed by proposed solutions or mitigations"
            }
        
        # Findings/results sections
        elif any(word in section_lower for word in ['finding', 'result', 'conclusion', 'outcome']):
            return {
                "objective": "Present the key findings, results, or conclusions from this section with supporting details and evidence.",
                "tone": "Factual and evidence-based",
                "length": "2-4 paragraphs, 200-300 words",
                "format": "Structured presentation of findings with supporting data or examples"
            }
        
        # Recommendation sections
        elif any(word in section_lower for word in ['recommendation', 'suggestion', 'next steps', 'action']):
            return {
                "objective": "Clearly present the recommendations, suggestions, or next steps outlined in this section.",
                "tone": "Actionable and directive",
                "length": "2-3 paragraphs, 150-250 words",
                "format": "Clear, actionable recommendations with rationale and expected outcomes"
            }
        
        # Default instructions for any other section
        else:
            return {
                "objective": f"Summarize and present the key information from the '{section_name}' section in a clear and organized manner.",
                "tone": "Professional and informative",
                "length": "2-3 paragraphs, 150-250 words",
                "format": "Well-structured paragraphs that clearly convey the main points and important details"
            }
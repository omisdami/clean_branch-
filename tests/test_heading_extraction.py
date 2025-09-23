import pytest
import tempfile
import os
from docx import Document
from fpdf import FPDF

from core.rag.ingestion.heading_extractor import HeadingExtractor

class TestHeadingExtraction:
    """Test dynamic heading extraction functionality"""
    
    def create_test_docx_with_headings(self) -> str:
        """Create a test DOCX file with various heading styles"""
        doc = Document()
        
        # Add different heading styles
        doc.add_heading('1. Introduction', level=1)
        doc.add_paragraph('This is the introduction content.')
        
        doc.add_heading('1.1 Background', level=2)
        doc.add_paragraph('Background information here.')
        
        doc.add_heading('2. Methodology', level=1)
        doc.add_paragraph('Methodology content.')
        
        doc.add_heading('2.1 Data Collection', level=2)
        doc.add_paragraph('Data collection details.')
        
        doc.add_heading('3. Findings', level=1)
        doc.add_paragraph('Key findings from the study.')
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            doc.save(tmp.name)
            return tmp.name
    
    def create_test_pdf_with_headings(self) -> str:
        """Create a test PDF file with headings"""
        pdf = FPDF()
        pdf.add_page()
        
        # Add headings with different font sizes
        pdf.set_font('Arial', 'B', 16)
        pdf.cell(0, 10, '1. Introduction', ln=True)
        pdf.set_font('Arial', '', 12)
        pdf.cell(0, 10, 'This is the introduction content.', ln=True)
        
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, '1.1 Background', ln=True)
        pdf.set_font('Arial', '', 12)
        pdf.cell(0, 10, 'Background information here.', ln=True)
        
        pdf.set_font('Arial', 'B', 16)
        pdf.cell(0, 10, '2. Methodology', ln=True)
        pdf.set_font('Arial', '', 12)
        pdf.cell(0, 10, 'Methodology content.', ln=True)
        
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
            pdf.output(tmp.name)
            return tmp.name
    
    def test_docx_heading_extraction(self):
        """Test DOCX heading extraction"""
        docx_path = self.create_test_docx_with_headings()
        
        try:
            extractor = HeadingExtractor()
            doc = Document(docx_path)
            headings = extractor.extract_docx_headings(doc)
            
            # Should detect multiple headings
            assert len(headings) >= 3
            
            # Check heading structure
            intro_heading = next((h for h in headings if 'Introduction' in h['heading']), None)
            assert intro_heading is not None
            assert intro_heading['level'] == 1
            assert len(intro_heading['content']) > 0
            
            # Check section IDs are unique
            section_ids = [h['section_id'] for h in headings]
            assert len(section_ids) == len(set(section_ids))
            
        finally:
            os.unlink(docx_path)
    
    def test_pdf_heading_extraction(self):
        """Test PDF heading extraction"""
        pdf_path = self.create_test_pdf_with_headings()
        
        try:
            extractor = HeadingExtractor()
            headings = extractor.extract_pdf_headings(pdf_path)
            
            # Should detect headings
            assert len(headings) >= 2
            
            # Check for numbered headings
            numbered_headings = [h for h in headings if h['heading'].startswith('1.') or h['heading'].startswith('2.')]
            assert len(numbered_headings) >= 2
            
        finally:
            os.unlink(pdf_path)
    
    def test_heading_mapping(self):
        """Test mapping detected headings to target sections"""
        extractor = HeadingExtractor()
        
        # Mock detected headings
        detected_headings = [
            {
                'heading': 'Executive Overview',
                'level': 1,
                'content': 'This is an executive overview...',
                'page_start': 1,
                'page_end': 1,
                'section_id': 'sec_001'
            },
            {
                'heading': 'Project Scope',
                'level': 1,
                'content': 'The scope includes...',
                'page_start': 2,
                'page_end': 2,
                'section_id': 'sec_002'
            },
            {
                'heading': 'Risk Analysis',
                'level': 1,
                'content': 'Key risks identified...',
                'page_start': 3,
                'page_end': 3,
                'section_id': 'sec_003'
            }
        ]
        
        target_sections = ['executive_summary', 'scope_of_work', 'risks']
        
        mapping = extractor.map_headings_to_sections(detected_headings, target_sections)
        
        # Should map executive overview to executive summary
        assert 'executive_summary' in mapping
        assert 'overview' in mapping['executive_summary']['heading'].lower()
        
        # Should map project scope to scope of work
        assert 'scope_of_work' in mapping
        assert 'scope' in mapping['scope_of_work']['heading'].lower()
        
        # Should map risk analysis to risks
        assert 'risks' in mapping
        assert 'risk' in mapping['risks']['heading'].lower()
    
    def test_fallback_mapping(self):
        """Test fallback when no headings match target sections"""
        extractor = HeadingExtractor()
        
        # Mock detected headings that don't match targets
        detected_headings = [
            {
                'heading': 'Random Section',
                'level': 1,
                'content': 'Some random content...',
                'page_start': 1,
                'page_end': 1,
                'section_id': 'sec_001'
            }
        ]
        
        target_sections = ['executive_summary']
        
        mapping = extractor.map_headings_to_sections(detected_headings, target_sections)
        
        # Should create fallback mapping
        assert 'executive_summary' in mapping
        assert 'Summarized' in mapping['executive_summary']['heading']
        assert len(mapping['executive_summary']['content']) > 0
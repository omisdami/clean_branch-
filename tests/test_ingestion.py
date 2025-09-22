import pytest
import tempfile
import os
from docx import Document
from fpdf import FPDF

from core.rag.ingestion.docx_extractor import DocxExtractor
from core.rag.ingestion.pdf_extractor import PdfExtractor
from core.rag.ingestion.extractor_factory import ExtractorFactory

class TestIngestion:
    """Test document ingestion functionality"""
    
    def create_test_docx(self) -> str:
        """Create a test DOCX file"""
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test paragraph.')
        doc.add_heading('Section 1', level=1)
        doc.add_paragraph('Content for section 1.')
        
        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = 'Header 1'
        table.cell(0, 1).text = 'Header 2'
        table.cell(1, 0).text = 'Data 1'
        table.cell(1, 1).text = 'Data 2'
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            doc.save(tmp.name)
            return tmp.name
    
    def create_test_pdf(self) -> str:
        """Create a test PDF file"""
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font('Arial', 'B', 16)
        pdf.cell(40, 10, 'Test Document')
        pdf.ln()
        pdf.set_font('Arial', '', 12)
        pdf.cell(40, 10, 'This is a test paragraph.')
        pdf.ln()
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(40, 10, 'Section 1')
        pdf.ln()
        pdf.set_font('Arial', '', 12)
        pdf.cell(40, 10, 'Content for section 1.')
        
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
            pdf.output(tmp.name)
            return tmp.name
    
    def test_docx_extraction(self):
        """Test DOCX document extraction"""
        docx_path = self.create_test_docx()
        
        try:
            extractor = DocxExtractor()
            document = extractor.extract(docx_path)
            
            # Verify metadata
            assert document.metadata.filetype == "docx"
            assert document.metadata.doc_id != ""
            assert document.metadata.pages >= 1
            
            # Verify sections
            assert len(document.sections) > 0
            assert any("Test Document" in section.title for section in document.sections)
            
            # Verify tables
            assert len(document.tables) > 0
            table = document.tables[0]
            assert "Header 1" in table.headers
            assert len(table.rows) > 0
            
        finally:
            os.unlink(docx_path)
    
    def test_pdf_extraction(self):
        """Test PDF document extraction"""
        pdf_path = self.create_test_pdf()
        
        try:
            extractor = PdfExtractor()
            document = extractor.extract(pdf_path)
            
            # Verify metadata
            assert document.metadata.filetype == "pdf"
            assert document.metadata.doc_id != ""
            assert document.metadata.pages >= 1
            
            # Verify sections
            assert len(document.sections) > 0
            
        finally:
            os.unlink(pdf_path)
    
    def test_extractor_factory(self):
        """Test extractor factory"""
        docx_path = self.create_test_docx()
        pdf_path = self.create_test_pdf()
        
        try:
            # Test DOCX
            docx_extractor = ExtractorFactory.get_extractor(docx_path)
            assert isinstance(docx_extractor, DocxExtractor)
            
            # Test PDF
            pdf_extractor = ExtractorFactory.get_extractor(pdf_path)
            assert isinstance(pdf_extractor, PdfExtractor)
            
            # Test unsupported
            unsupported_extractor = ExtractorFactory.get_extractor("test.txt")
            assert unsupported_extractor is None
            
        finally:
            os.unlink(docx_path)
            os.unlink(pdf_path)
    
    def test_supported_filetypes(self):
        """Test supported filetypes"""
        supported = ExtractorFactory.supported_filetypes()
        assert 'docx' in supported
        assert 'pdf' in supported